"""원본 문서(docx/hwpx)의 구조(표·스타일·레이아웃)를 보존하면서
문단 텍스트만 교정 결과로 교체하고, 수정된 부분을 빨간 글자로 표시한다.

사용 흐름:
  1) collect_*()  : 원본에서 교정 대상 문단 텍스트 목록 추출 (구조 정보 유지)
  2) (앱에서 Claude로 문단별 교정)
  3) apply_*()    : 교정문을 원본 구조 안에 diff 기반 빨간 run으로 되써넣기
"""

import io
import re
import zipfile
from xml.etree import ElementTree

from export import diff_segments

RED_RGB = (0xC0, 0x00, 0x00)

# ── DOCX ──────────────────────────────────────────────────────


def _iter_docx_paragraphs(container):
    """본문·표(중첩 표 포함)의 문단 순회"""

    def walk_cell(cell):
        for p in cell.paragraphs:
            yield p
        for t in cell.tables:
            for row in t.rows:
                for c in row.cells:
                    yield from walk_cell(c)

    for p in container.paragraphs:
        yield p
    for t in container.tables:
        for row in t.rows:
            for c in row.cells:
                yield from walk_cell(c)


def collect_docx(file_like):
    """(Document 객체, 문단 객체 목록, 문단 텍스트 목록) 반환"""
    from docx import Document
    doc = Document(file_like)
    paras = [p for p in _iter_docx_paragraphs(doc) if p.text.strip()]
    return doc, paras, [p.text for p in paras]


def apply_docx(doc, paras, corrected_texts):
    """교정문을 문단별로 되써넣고 수정 부분을 빨간 글자로. (docx bytes, 변경문단수) 반환"""
    from docx.shared import RGBColor
    red = RGBColor(*RED_RGB)
    n_changed = 0
    for para, corrected in zip(paras, corrected_texts):
        original = para.text
        corrected = (corrected or "").replace("\n", " ").strip()
        if not corrected or corrected == original.strip():
            continue
        # 원래 서체 특성(첫 run 기준) 보존
        base_font = para.runs[0].font if para.runs else None
        base = {}
        if base_font is not None:
            base = {"name": base_font.name, "size": base_font.size,
                    "bold": base_font.bold, "italic": base_font.italic}
            # 한글 글꼴(eastAsia)도 보존
            rPr = para.runs[0]._element.rPr
            ea = None
            if rPr is not None and rPr.rFonts is not None:
                ea = rPr.rFonts.get(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia")
            base["east_asia"] = ea
        for r in list(para.runs):
            r._element.getparent().remove(r._element)
        for text, changed in diff_segments(original, corrected):
            run = para.add_run(text.replace("\n", " "))
            if base.get("name"):
                run.font.name = base["name"]
            if base.get("size"):
                run.font.size = base["size"]
            if base.get("bold") is not None:
                run.font.bold = base["bold"]
            if base.get("italic") is not None:
                run.font.italic = base["italic"]
            if base.get("east_asia"):
                run._element.get_or_add_rPr().get_or_add_rFonts().set(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
                    base["east_asia"])
            if changed:
                run.font.color.rgb = red
        n_changed += 1
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), n_changed


# ── HWPX ──────────────────────────────────────────────────────

_HWPX_NS = {
    "hs": "http://www.hancom.co.kr/hwpml/2011/section",
    "hp": "http://www.hancom.co.kr/hwpml/2011/paragraph",
    "hp10": "http://www.hancom.co.kr/hwpml/2016/paragraph",
    "ha": "http://www.hancom.co.kr/hwpml/2011/app",
    "hc": "http://www.hancom.co.kr/hwpml/2011/core",
    "hh": "http://www.hancom.co.kr/hwpml/2011/head",
    "hhs": "http://www.hancom.co.kr/hwpml/2011/history",
    "hm": "http://www.hancom.co.kr/hwpml/2011/master-page",
    "hpf": "http://www.hancom.co.kr/schema/2011/hpf",
    "dc": "http://purl.org/dc/elements/1.1/",
    "opf": "http://www.idpf.org/2007/opf/",
    "ooxmlchart": "http://www.hancom.co.kr/hwpml/2016/ooxmlchart",
    "hwpunitchar": "http://www.hancom.co.kr/hwpml/2016/HwpUnitChar",
    "epub": "http://www.idpf.org/2007/ops",
    "config": "urn:oasis:names:tc:opendocument:xmlns:config:1.0",
}
for _pfx, _uri in _HWPX_NS.items():
    ElementTree.register_namespace(_pfx, _uri)

_HP = "{%s}" % _HWPX_NS["hp"]
_HH = "{%s}" % _HWPX_NS["hh"]


def _hwpx_para_text_runs(p_el):
    """문단의 '직계 텍스트 run'들만 (표 등 객체 run 제외).
    반환: [(run_el, run_text)], 문단 전체 텍스트"""
    text_runs, full = [], []
    for run in p_el.findall(f"{_HP}run"):
        has_obj = any(ch.tag != f"{_HP}t" for ch in run)
        ts = run.findall(f"{_HP}t")
        if has_obj or not ts:
            continue  # 표·그림·개체 run은 건드리지 않음
        rt = "".join((t.text or "") + "".join(
            (sub.tail or "") for sub in t) for t in ts)
        text_runs.append((run, rt))
        full.append(rt)
    return text_runs, "".join(full)


def collect_hwpx(file_bytes):
    """(state, 문단 텍스트 목록) 반환. state는 apply_hwpx에 그대로 전달."""
    zf = zipfile.ZipFile(io.BytesIO(file_bytes))
    section_names = sorted(
        n for n in zf.namelist() if re.match(r"Contents/section\d+\.xml$", n))
    if not section_names:
        raise ValueError("HWPX 본문(section XML)을 찾을 수 없어요.")
    trees, para_refs, texts = {}, [], []
    for name in section_names:
        root = ElementTree.fromstring(zf.read(name))
        trees[name] = root
        for p_el in root.iter(f"{_HP}p"):
            runs, full = _hwpx_para_text_runs(p_el)
            if full.strip():
                para_refs.append((name, p_el, runs))
                texts.append(full)
    state = {"zip_bytes": file_bytes, "trees": trees,
             "sections": section_names, "para_refs": para_refs}
    return state, texts


def _hwpx_add_red_charpr(header_xml):
    """header.xml에 빨간 charPr 추가(기본 charPr 복제, 연속 id). (새 xml, red_id) 반환"""
    # charPr는 자기닫힘(<hh:charPr .../>)과 일반(<hh:charPr>...</hh:charPr>) 두 형태 모두 존재
    m = re.search(r'<hh:charPr id="0"[^>]*/>', header_xml)
    if not m:
        m = re.search(r'<hh:charPr id="0".*?</hh:charPr>', header_xml, re.S)
    if not m:
        raise ValueError("문서에서 기본 글자 속성을 찾지 못했어요.")
    ids = [int(x) for x in re.findall(r'<hh:charPr id="(\d+)"', header_xml)]
    red_id = str(max(ids) + 1)
    block = m.group(0).replace('id="0"', f'id="{red_id}"', 1)
    if 'textColor="' in block:
        block = re.sub(r'textColor="#[0-9A-Fa-f]{6}"', 'textColor="#C00000"', block, count=1)
    else:
        block = block.replace(f'id="{red_id}"', f'id="{red_id}" textColor="#C00000"', 1)
    header_xml = header_xml.replace("</hh:charProperties>", block + "</hh:charProperties>", 1)
    header_xml = re.sub(r'(<hh:charProperties itemCnt=")(\d+)(")',
                        lambda mm: mm.group(1) + str(int(mm.group(2)) + 1) + mm.group(3),
                        header_xml, count=1)
    return header_xml, red_id


def apply_hwpx(state, corrected_texts):
    """교정문을 원본 hwpx 구조 안에 되써넣기. (hwpx bytes, 변경문단수) 반환"""
    src = zipfile.ZipFile(io.BytesIO(state["zip_bytes"]))
    header_xml, red_id = _hwpx_add_red_charpr(
        src.read("Contents/header.xml").decode("utf-8"))

    n_changed = 0
    for (sec_name, p_el, runs), corrected in zip(state["para_refs"], corrected_texts):
        original = "".join(rt for _, rt in runs)
        corrected = (corrected or "").replace("\n", " ").strip()
        if not corrected or corrected == original.strip():
            continue
        base_charpr = runs[0][0].get("charPrIDRef", "0")
        children = list(p_el)
        insert_at = children.index(runs[0][0])
        for run_el, _ in runs:
            p_el.remove(run_el)
        for text, changed in diff_segments(original, corrected):
            run_el = ElementTree.Element(f"{_HP}run")
            run_el.set("charPrIDRef", red_id if changed else base_charpr)
            t_el = ElementTree.SubElement(run_el, f"{_HP}t")
            t_el.text = text.replace("\n", " ")
            p_el.insert(insert_at, run_el)
            insert_at += 1
        n_changed += 1

    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("mimetype", src.read("mimetype"),
                    compress_type=zipfile.ZIP_STORED)
        for name in src.namelist():
            if name == "mimetype":
                continue
            elif name == "Contents/header.xml":
                zf.writestr(name, header_xml)
            elif name in state["trees"]:
                xml = ElementTree.tostring(
                    state["trees"][name], encoding="unicode")
                zf.writestr(name, '<?xml version="1.0" encoding="UTF-8" standalone="yes" ?>' + xml)
            else:
                zf.writestr(name, src.read(name))
    src.close()
    return out.getvalue(), n_changed


# ── 문단 번호 매핑 (Claude 교정 왕복용) ──────────────────────

def chunk_numbered(texts, limit=3000):
    """문단 목록 → [(시작번호, '⟦N⟧ 텍스트' 블록)] 구간 목록"""
    chunks, cur, cur_len, start = [], [], 0, 0
    for i, t in enumerate(texts):
        line = f"⟦{i + 1}⟧ {t}"
        if cur and cur_len + len(line) > limit:
            chunks.append((start, "\n".join(cur)))
            cur, cur_len, start = [], 0, i
        cur.append(line)
        cur_len += len(line)
    if cur:
        chunks.append((start, "\n".join(cur)))
    return chunks


def parse_numbered(response, texts):
    """Claude 응답에서 ⟦N⟧ 문단 파싱 → {번호(0기준): 교정문}. 누락 번호는 무시(원문 유지)."""
    result = {}
    body = response.split("[수정 설명]")[0]
    for m in re.finditer(r"⟦(\d+)⟧\s*(.*?)(?=⟦\d+⟧|$)", body, re.S):
        idx = int(m.group(1)) - 1
        if 0 <= idx < len(texts):
            txt = m.group(2).strip()
            if txt:
                result[idx] = txt
    explanation = ""
    if "[수정 설명]" in response:
        explanation = response.split("[수정 설명]", 1)[1].strip()
    return result, explanation
