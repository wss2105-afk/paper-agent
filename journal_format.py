"""학술지 투고 형식 맞추기 — 형식 소스(웹/파일) 수집 + 규정에 맞춘 문서 생성.

- fetch_url_text: 학술지 author guidelines 웹페이지에서 텍스트 추출
- build_docx_from_markdown: 재편집된 원고(마크다운)를 레이아웃 적용해 docx 생성
- build_hwpx_from_markdown: 같은 내용을 hwpx로 생성(내용 위주, 레이아웃 제약)
LLM(규정 요약·원고 재편집)은 app.py에서 담당하고, 이 모듈은 수집·생성만 한다.
"""

import io
import re

_UA = ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
       "(KHTML, like Gecko) Chrome/120.0 Safari/537.36")

DEFAULT_LAYOUT = {
    "font": "",                # 한글 글꼴명 (빈값이면 기본)
    "size_pt": 11,
    "line_spacing": 1.6,
    "margin_cm": {"top": 2.5, "bottom": 2.5, "left": 2.5, "right": 2.5},
}


def fetch_url_text(url, max_chars=16000):
    """학술지 규정 웹페이지 → 본문 텍스트. requests로 서버측 fetch 후 태그 제거."""
    import requests

    if not re.match(r"^https?://", url, re.I):
        url = "https://" + url.strip()
    resp = requests.get(url, headers={"User-Agent": _UA}, timeout=20)
    resp.raise_for_status()
    if resp.encoding is None or resp.encoding.lower() == "iso-8859-1":
        resp.encoding = resp.apparent_encoding
    html = resp.text
    # 스크립트/스타일/주석 제거
    html = re.sub(r"(?is)<(script|style|noscript|head)[^>]*>.*?</\1>", " ", html)
    html = re.sub(r"(?s)<!--.*?-->", " ", html)
    # 블록 태그는 줄바꿈으로
    html = re.sub(r"(?i)<(br|/p|/div|/li|/tr|/h[1-6])\s*/?>", "\n", html)
    text = re.sub(r"(?s)<[^>]+>", " ", html)
    # 엔티티 최소 복원
    for a, b in [("&nbsp;", " "), ("&amp;", "&"), ("&lt;", "<"),
                 ("&gt;", ">"), ("&quot;", '"'), ("&#39;", "'")]:
        text = text.replace(a, b)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n[ \t]*\n\s*", "\n\n", text)
    text = text.strip()
    if not text:
        raise ValueError("페이지에서 텍스트를 추출하지 못했어요. (로그인/자바스크립트 전용 페이지일 수 있어요)")
    return text[:max_chars]


def _set_korean_font(style, font_name):
    """python-docx 스타일에 한글(eastAsia) 글꼴까지 지정."""
    from docx.oxml.ns import qn
    style.font.name = font_name
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        rfonts.set(qn(attr), font_name)


def build_docx_from_markdown(md_text, layout=None, title=""):
    """재편집된 원고(간단 마크다운: # / ## / ### 제목, 나머지 문단)를
    학술지 레이아웃(여백·글꼴·크기·줄간격) 적용해 docx로 생성."""
    from docx import Document
    from docx.shared import Pt, Cm

    lay = {**DEFAULT_LAYOUT, **(layout or {})}
    doc = Document()

    sec = doc.sections[0]
    mc = {**DEFAULT_LAYOUT["margin_cm"], **(lay.get("margin_cm") or {})}
    sec.top_margin = Cm(_num(mc.get("top"), 2.5))
    sec.bottom_margin = Cm(_num(mc.get("bottom"), 2.5))
    sec.left_margin = Cm(_num(mc.get("left"), 2.5))
    sec.right_margin = Cm(_num(mc.get("right"), 2.5))

    normal = doc.styles["Normal"]
    if lay.get("font"):
        _set_korean_font(normal, lay["font"])
    normal.font.size = Pt(_num(lay.get("size_pt"), 11))
    normal.paragraph_format.line_spacing = _num(lay.get("line_spacing"), 1.6)

    if title:
        h = doc.add_heading(title, level=0)

    def _inline(s):
        """마크다운 강조 마커 제거 (**굵게**, *기울임*) — Word엔 텍스트만."""
        s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
        s = re.sub(r"(?<![\w*])\*([^*\n]+?)\*(?![\w*])", r"\1", s)
        return s

    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        if line.lstrip().startswith("|"):
            # 마크다운 표 → 실제 Word 표 (구분선 행 |---|---| 은 건너뜀)
            rows = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                cells = [_inline(c.strip())
                         for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-+:?", c) for c in cells if c):
                    rows.append(cells)
                i += 1
            if rows:
                ncol = max(len(r) for r in rows)
                tbl = doc.add_table(rows=len(rows), cols=ncol)
                try:
                    tbl.style = "Table Grid"
                except Exception:
                    pass
                for ri, r in enumerate(rows):
                    for ci in range(ncol):
                        tbl.cell(ri, ci).text = r[ci] if ci < len(r) else ""
            continue
        if line.startswith("### "):
            doc.add_heading(_inline(line[4:].strip()), level=3)
        elif line.startswith("## "):
            doc.add_heading(_inline(line[3:].strip()), level=2)
        elif line.startswith("# "):
            doc.add_heading(_inline(line[2:].strip()), level=1)
        else:
            doc.add_paragraph(_inline(line))
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_hwpx_from_markdown(md_text, title=""):
    """같은 내용을 hwpx로 생성 (한글 템플릿 기반, 내용 위주)."""
    from export import _hwpx_template_bytes, _hwpx_paragraph
    import zipfile

    src = zipfile.ZipFile(io.BytesIO(_hwpx_template_bytes()))
    section = src.read("Contents/section0.xml").decode("utf-8")

    pid = 80000000
    body = ""
    if title:
        body += _hwpx_paragraph([(title, False)], pid, "0")
        pid += 1
    for raw in md_text.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        # 마크다운 heading 기호는 제거하고 텍스트만 (hwpx는 문단으로)
        line = re.sub(r"^#{1,3}\s+", "", line)
        pid += 1
        body += _hwpx_paragraph([(line, False)], pid, "0")

    section = section.replace("</hs:sec>", body + "</hs:sec>", 1)
    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("mimetype", src.read("mimetype"),
                    compress_type=zipfile.ZIP_STORED)
        for name in src.namelist():
            if name == "mimetype":
                continue
            elif name == "Contents/section0.xml":
                zf.writestr(name, section)
            else:
                zf.writestr(name, src.read(name))
    src.close()
    return out.getvalue()


def _num(v, default):
    try:
        f = float(v)
        return f if f > 0 else default
    except (TypeError, ValueError):
        return default
