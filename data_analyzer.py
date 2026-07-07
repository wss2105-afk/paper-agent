import io
import re
import zipfile
from xml.etree import ElementTree

import pandas as pd


def extract_hwpx_text(file_like):
    """HWPX(한글 표준 XML 포맷) 본문 텍스트 추출.
    HWPX는 zip 압축 안에 Contents/section*.xml 로 본문을 담는다.
    문서 순서대로 순회하며 문단(<hp:p>) 단위로 텍스트(<hp:t>)를 모은다."""
    try:
        zf = zipfile.ZipFile(file_like)
    except zipfile.BadZipFile:
        raise ValueError(
            "HWPX 형식이 아니에요. 구형 .hwp 파일이면 한글에서 "
            "'다른 이름으로 저장 → HWPX'로 저장한 뒤 올려주세요."
        )
    with zf:
        sections = sorted(
            n for n in zf.namelist()
            if re.match(r"Contents/section\d+\.xml$", n)
        )
        if not sections:
            raise ValueError("HWPX 본문(section XML)을 찾을 수 없어요. 파일이 손상됐을 수 있어요.")
        paras = [""]
        for sec_name in sections:
            root = ElementTree.fromstring(zf.read(sec_name))
            for el in root.iter():
                tag = el.tag.split("}")[-1]
                if tag == "p":  # 새 문단 (표 안 문단 포함, 문서 순서 유지)
                    if paras[-1].strip():
                        paras.append("")
                elif tag == "t" and el.text:
                    paras[-1] += el.text
    return "\n".join(p for p in paras if p.strip())


# HWP 5.0 텍스트 레코드 안의 제어문자 처리:
#   인라인/확장 제어(1~23 중 아래 목록)는 자기 자신 포함 8 WCHAR(16바이트)를 차지
#   나머지(0, 10, 13, 24~31)는 1 WCHAR — 13(문단끝)/10(줄바꿈)만 개행으로 반영
_HWP_CTRL_8WCHAR = {1, 2, 3, 4, 5, 6, 7, 8, 9, 11, 12, 14, 15, 16, 17, 18, 19, 20, 21, 22, 23}


def extract_hwp_text(file_like):
    """구형 한글(.hwp, HWP 5.0 바이너리) 본문 텍스트 추출.
    BodyText 섹션의 레코드를 파싱해 문단 텍스트(HWPTAG_PARA_TEXT=67)를 모은다.
    본문 파싱이 실패하면 미리보기(PrvText) 스트림으로 폴백."""
    import struct
    import zlib

    import olefile

    if not olefile.isOleFile(file_like):
        raise ValueError("HWP 5.0 형식이 아니에요. (한글 2002 이전 파일이거나 손상됐을 수 있어요)")
    ole = olefile.OleFileIO(file_like)
    try:
        header = ole.openstream("FileHeader").read()
        if not header.startswith(b"HWP Document File"):
            raise ValueError("HWP 파일 헤더를 인식할 수 없어요.")
        flags = struct.unpack("<I", header[36:40])[0]
        if flags & 0x2:
            raise ValueError("암호가 걸린 .hwp 파일이에요. 한글에서 암호를 해제한 뒤 올려주세요.")
        compressed = bool(flags & 0x1)

        def parse_bodytext():
            secs = sorted(
                (e for e in ole.listdir() if len(e) == 2 and e[0] == "BodyText"),
                key=lambda e: int(e[1].replace("Section", "") or 0),
            )
            if not secs:
                raise ValueError("본문(BodyText)이 없어요.")
            paras = []
            for entry in secs:
                data = ole.openstream(entry).read()
                if compressed:
                    data = zlib.decompress(data, -15)
                i = 0
                while i + 4 <= len(data):
                    hdr = struct.unpack("<I", data[i:i + 4])[0]
                    tag = hdr & 0x3FF
                    size = (hdr >> 20) & 0xFFF
                    i += 4
                    if size == 0xFFF:  # 확장 크기
                        size = struct.unpack("<I", data[i:i + 4])[0]
                        i += 4
                    if tag == 67:  # HWPTAG_PARA_TEXT
                        payload = data[i:i + size]
                        j, buf = 0, []
                        while j + 2 <= len(payload):
                            code = struct.unpack("<H", payload[j:j + 2])[0]
                            if code in _HWP_CTRL_8WCHAR:
                                j += 16
                            elif code < 32:
                                if code in (10, 13):
                                    buf.append("\n")
                                j += 2
                            else:
                                buf.append(chr(code))
                                j += 2
                        text = "".join(buf).strip()
                        if text:
                            paras.append(text)
                    i += size
            return "\n".join(paras)

        try:
            text = parse_bodytext()
            if text.strip():
                return text
        except ValueError:
            raise
        except Exception:
            pass  # 본문 파싱 실패 → 미리보기 폴백

        # 폴백: PrvText(미리보기, 앞부분 일부만 담김)
        if ole.exists("PrvText"):
            prv = ole.openstream("PrvText").read().decode("utf-16-le", errors="ignore").strip("\x00 \r\n")
            if prv.strip():
                return prv + "\n\n[안내] 본문 전체 추출에 실패해 미리보기 텍스트만 불러왔어요. 한글에서 HWPX로 저장 후 올리면 전체가 추출됩니다."
        raise ValueError("본문 텍스트를 추출하지 못했어요. 한글에서 '다른 이름으로 저장 → HWPX'로 저장 후 올려주세요.")
    finally:
        ole.close()


def load_codebook_text(uploaded_file):
    """코딩북(변수 설명서) 파일에서 텍스트 추출.
    Excel/CSV는 표 그대로, Word는 문단+표, 한글/PDF/텍스트도 지원."""
    name = uploaded_file.name.lower()
    if name.endswith((".xlsx", ".xls")):
        df = pd.read_excel(uploaded_file)
        return df.to_string(index=False)
    if name.endswith(".csv"):
        df = pd.read_csv(uploaded_file)
        return df.to_string(index=False)
    if name.endswith(".docx"):
        from docx import Document
        doc = Document(uploaded_file)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for t in doc.tables:  # 코딩북은 표 형태가 많음
            for row in t.rows:
                parts.append(" | ".join(c.text.strip() for c in row.cells))
        return "\n".join(parts)
    if name.endswith(".hwpx"):
        return extract_hwpx_text(uploaded_file)
    if name.endswith(".hwp"):
        return extract_hwp_text(uploaded_file)
    if name.endswith(".pdf"):
        import pdfplumber
        text = ""
        with pdfplumber.open(uploaded_file) as pdf:
            for page in pdf.pages[:30]:
                t = page.extract_text()
                if t:
                    text += t + "\n"
        return text
    return uploaded_file.read().decode("utf-8", errors="ignore")


def load_file(uploaded_file):
    """Excel, SPSS, CSV, 텍스트, Word, 한글(HWP/HWPX) 파일 로드"""
    name = uploaded_file.name.lower()

    if name.endswith(".sav"):
        import pyreadstat
        with io.BytesIO(uploaded_file.read()) as buf:
            df, meta = pyreadstat.read_sav(buf)
        return "quantitative", df, meta

    elif name.endswith((".xlsx", ".xls")):
        df = pd.read_excel(uploaded_file)
        return "quantitative", df, None

    elif name.endswith(".csv"):
        df = pd.read_csv(uploaded_file)
        return "quantitative", df, None

    elif name.endswith(".txt"):
        text = uploaded_file.read().decode("utf-8", errors="ignore")
        return "qualitative", text, None

    elif name.endswith(".docx"):
        from docx import Document
        doc = Document(uploaded_file)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return "qualitative", text, None

    elif name.endswith(".hwpx"):
        text = extract_hwpx_text(uploaded_file)
        return "qualitative", text, None

    elif name.endswith(".hwp"):
        text = extract_hwp_text(uploaded_file)
        return "qualitative", text, None

    else:
        raise ValueError(
            "지원 형식: Excel(.xlsx/.xls), SPSS(.sav), CSV(.csv), "
            "텍스트(.txt), Word(.docx), 한글(.hwpx)"
        )


def summarize_dataframe(df, meta=None):
    """정량 데이터 구조 요약 (Claude에게 전달용)"""
    lines = []
    lines.append(f"- 총 행(응답자) 수: {len(df)}")
    lines.append(f"- 총 열(변수) 수: {len(df.columns)}")
    lines.append("")
    lines.append("## 변수 목록")

    for col in df.columns:
        dtype = df[col].dtype
        missing = df[col].isnull().sum()
        missing_pct = missing / len(df) * 100

        label = ""
        if meta and hasattr(meta, "column_labels") and meta.column_labels:
            label_val = meta.column_labels.get(col, "")
            if label_val:
                label = f" [{label_val}]"

        if pd.api.types.is_numeric_dtype(dtype):
            mean = df[col].mean()
            std = df[col].std()
            min_v = df[col].min()
            max_v = df[col].max()
            lines.append(
                f"- **{col}**{label}: 수치형 | 평균={mean:.2f}, SD={std:.2f}, "
                f"범위={min_v}~{max_v}, 결측={missing}({missing_pct:.1f}%)"
            )
        else:
            unique = df[col].nunique()
            top_vals = df[col].value_counts().head(5).to_dict()
            lines.append(
                f"- **{col}**{label}: 범주형 | 고유값={unique}개, "
                f"주요값={top_vals}, 결측={missing}({missing_pct:.1f}%)"
            )

    return "\n".join(lines)


def summarize_interview(text):
    """인터뷰 텍스트 구조 요약"""
    lines = text.splitlines()
    non_empty = [l for l in lines if l.strip()]
    word_count = len(text.split())
    char_count = len(text)

    # 간단한 참여자 구분 시도 (Q:, A:, 면담자, 참여자 등 패턴)
    speakers = set()
    for line in non_empty:
        for prefix in ["Q:", "A:", "면담자:", "참여자:", "연구자:", "교사:", "학생:"]:
            if line.strip().startswith(prefix):
                speakers.add(prefix.rstrip(":"))

    summary = [
        f"- 총 글자 수: {char_count:,}",
        f"- 총 단어 수: {word_count:,}",
        f"- 총 줄 수: {len(non_empty)}",
    ]
    if speakers:
        summary.append(f"- 감지된 화자: {', '.join(sorted(speakers))}")
    summary.append("")
    summary.append("## 텍스트 앞부분 미리보기")
    summary.append(text[:800] + ("..." if len(text) > 800 else ""))

    return "\n".join(summary)


def get_preview(df, n=5):
    return df.head(n)


def get_basic_stats(df):
    numeric_cols = df.select_dtypes(include="number")
    if numeric_cols.empty:
        return None
    return numeric_cols.describe().round(2)
