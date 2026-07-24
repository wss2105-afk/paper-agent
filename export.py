import difflib
import io
import re

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def to_word(title: str, content: str, topic: str = "") -> bytes:
    doc = Document()

    # 제목
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if topic:
        sub = doc.add_paragraph(f"주제: {topic}")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # 본문 — "**참고문헌:**" 기준으로 본문/참고문헌 분리
    if "**참고문헌:**" in content:
        body, refs = content.split("**참고문헌:**", 1)
    elif "참고문헌:" in content:
        body, refs = content.split("참고문헌:", 1)
    else:
        body, refs = content, ""

    # 본문 단락
    for line in body.strip().splitlines():
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph(line)
        p.style.font.size = Pt(11)

    # 참고문헌 섹션
    if refs.strip():
        doc.add_paragraph()
        ref_heading = doc.add_heading("참고문헌", level=2)
        for line in refs.strip().splitlines():
            line = line.strip().lstrip("-").strip()
            if line:
                doc.add_paragraph(line, style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def diff_segments(original: str, corrected: str):
    """원문 대비 교정문의 단어 단위 diff.
    교정문 기준 [(텍스트, 변경여부)] 목록 반환 — 변경/추가된 조각만 True."""
    o = re.split(r"(\s+)", original)
    c = re.split(r"(\s+)", corrected)
    sm = difflib.SequenceMatcher(None, o, c, autojunk=False)
    segs = []
    for tag, _i1, _i2, j1, j2 in sm.get_opcodes():
        chunk = "".join(c[j1:j2])
        if not chunk:
            continue
        # 공백만 바뀐 조각은 빨간 표시 대상에서 제외
        changed = tag != "equal" and bool(chunk.strip())
        segs.append((chunk, changed))
    return segs


def to_word_redline(original: str, corrected: str, explanation: str = "",
                    title: str = "교정 결과") -> bytes:
    """교정문을 Word로 내보내되, 원문에서 수정된 부분을 빨간 글자로 표시."""
    doc = Document()
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note = doc.add_paragraph("※ 빨간색 글자 = 원문에서 수정·추가된 부분")
    note.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    note.runs[0].font.size = Pt(9)
    doc.add_paragraph()

    para = doc.add_paragraph()
    for text, changed in diff_segments(original, corrected):
        parts = text.split("\n")
        for k, part in enumerate(parts):
            if k > 0:
                para = doc.add_paragraph()
            if part:
                run = para.add_run(part)
                run.font.size = Pt(11)
                if changed:
                    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    if explanation.strip():
        doc.add_paragraph()
        doc.add_heading("수정 설명", level=2)
        for line in explanation.strip().splitlines():
            line = line.strip()
            if line:
                doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── HWPX(한글) 내보내기 ───────────────────────────────────────
# assets/hwpx_template.hwpx: 한글이 만든 빈 문서(hwpxlib 프로젝트, Apache-2.0).
# 유효한 실제 파일을 템플릿으로 쓰고 본문/글자속성만 주입해 호환성을 확보한다.

_HWPX_TEMPLATE = None


def _hwpx_template_bytes():
    global _HWPX_TEMPLATE
    if _HWPX_TEMPLATE is None:
        from pathlib import Path
        _HWPX_TEMPLATE = (Path(__file__).parent / "assets" / "hwpx_template.hwpx").read_bytes()
    return _HWPX_TEMPLATE


def _hwpx_escape(text: str) -> str:
    return (text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            .replace('"', "&quot;"))


def _hwpx_paragraph(runs, para_id, red_id) -> str:
    """runs: [(텍스트, 빨간여부)] → <hp:p> XML"""
    run_xml = "".join(
        f'<hp:run charPrIDRef="{red_id if red else "0"}">'
        f"<hp:t>{_hwpx_escape(t)}</hp:t></hp:run>"
        for t, red in runs if t
    ) or '<hp:run charPrIDRef="0"><hp:t></hp:t></hp:run>'
    return (f'<hp:p id="{para_id}" paraPrIDRef="3" styleIDRef="0" '
            f'pageBreak="0" columnBreak="0" merged="0">{run_xml}</hp:p>')


def to_hwpx_redline(original: str, corrected: str, explanation: str = "") -> bytes:
    """교정문을 한글(.hwpx)로 내보내되, 수정된 부분을 빨간 글자로 표시."""
    import zipfile

    src = zipfile.ZipFile(io.BytesIO(_hwpx_template_bytes()))
    header = src.read("Contents/header.xml").decode("utf-8")
    section = src.read("Contents/section0.xml").decode("utf-8")

    # 1) 빨간 글자 charPr 추가: 기본 charPr(id=0)을 복제해 색만 변경.
    #    한글이 charPrIDRef를 순번으로 해석하는 경우가 있어, id는 반드시
    #    기존 id들에 이어지는 연속 번호로 부여한다 (id 목록과 순번 일치 유지).
    m = re.search(r'<hh:charPr id="0".*?</hh:charPr>', header, re.S)
    if not m:
        raise ValueError("HWPX 템플릿에서 기본 글자 속성을 찾지 못했어요.")
    existing_ids = [int(x) for x in re.findall(r'<hh:charPr id="(\d+)"', header)]
    red_id = str(max(existing_ids) + 1)
    red_charpr = (m.group(0)
                  .replace('id="0"', f'id="{red_id}"', 1)
                  .replace('textColor="#000000"', 'textColor="#C00000"', 1))
    header = header.replace("</hh:charProperties>", red_charpr + "</hh:charProperties>", 1)
    header = re.sub(r'(<hh:charProperties itemCnt=")(\d+)(")',
                    lambda mm: mm.group(1) + str(int(mm.group(2)) + 1) + mm.group(3),
                    header, count=1)

    # 2) 본문 문단 구성 (문단 경계 = 개행)
    paras, cur = [], []
    for text, changed in diff_segments(original, corrected):
        parts = text.split("\n")
        for k, part in enumerate(parts):
            if k > 0:
                paras.append(cur)
                cur = []
            if part:
                cur.append((part, changed))
    paras.append(cur)

    pid = 90000000
    body_xml = _hwpx_paragraph([("※ 빨간색 글자 = 원문에서 수정·추가된 부분", False)], pid, red_id)
    for p_runs in paras:
        pid += 1
        body_xml += _hwpx_paragraph(p_runs, pid, red_id)
    if explanation.strip():
        pid += 1
        body_xml += _hwpx_paragraph([("", False)], pid, red_id)
        pid += 1
        body_xml += _hwpx_paragraph([("[수정 설명]", False)], pid, red_id)
        for line in explanation.strip().splitlines():
            if line.strip():
                pid += 1
                body_xml += _hwpx_paragraph([(line.strip(), False)], pid, red_id)

    section = section.replace("</hs:sec>", body_xml + "</hs:sec>", 1)

    # 3) zip 재조립 (mimetype은 관례대로 무압축 선두 배치)
    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("mimetype", src.read("mimetype"), compress_type=zipfile.ZIP_STORED)
        for name in src.namelist():
            if name == "mimetype":
                continue
            elif name == "Contents/header.xml":
                zf.writestr(name, header)
            elif name == "Contents/section0.xml":
                zf.writestr(name, section)
            else:
                zf.writestr(name, src.read(name))
    src.close()
    return out.getvalue()


def _fmt_cell(v) -> str:
    """표 셀 값 → 문자열 (float는 유효숫자 4자리, NaN은 빈칸)"""
    import math
    if v is None:
        return ""
    if isinstance(v, float):
        if math.isnan(v):
            return ""
        s = f"{v:.4f}".rstrip("0").rstrip(".")
        return s if s not in ("", "-") else "0"
    return str(v)


def to_stats_docx(title: str, tables: dict, summary: str = "",
                  interp: str = "", note: str = "") -> bytes:
    """통계 분석 결과(결과표 + 요약 + 논문용 해석)를 Word(.docx)로 내보낸다.
    표는 실제 Word 표(Table Grid)로 만들어 한글(HWP)에서도 그대로 열린다."""
    import pandas as pd
    from docx.oxml.ns import qn

    doc = Document()
    # 한글(HWP) 호환을 위해 기본 글꼴에 한글 글꼴(eastAsia)을 명시
    normal = doc.styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal.font.size = Pt(10)
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "맑은 고딕")

    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if note:
        sub = doc.add_paragraph(note)
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        sub.runs[0].font.size = Pt(9)
    doc.add_paragraph()

    for tname, tdf in (tables or {}).items():
        doc.add_heading(str(tname), level=2)
        df = tdf
        include_index = bool(df.index.name) or not isinstance(df.index, pd.RangeIndex)
        cols = ([str(df.index.name or "")] if include_index else []) + [str(c) for c in df.columns]
        table = doc.add_table(rows=1, cols=max(1, len(cols)))
        table.style = "Table Grid"
        for j, c in enumerate(cols):
            cell = table.rows[0].cells[j]
            cell.text = c
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = True
        for i in range(len(df)):
            row = table.add_row().cells
            vals = ([df.index[i]] if include_index else []) + list(df.iloc[i])
            for j, v in enumerate(vals):
                row[j].text = _fmt_cell(v)
        doc.add_paragraph()

    if summary.strip():
        doc.add_heading("결과 요약", level=2)
        for line in summary.strip().splitlines():
            if line.strip():
                doc.add_paragraph(line.rstrip())

    if interp and interp.strip():
        doc.add_heading("논문용 결과 해석", level=2)
        for line in interp.strip().splitlines():
            line = line.strip().replace("**", "")
            if line:
                doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def to_markdown(title: str, content: str, topic: str = "") -> str:
    lines = [f"# {title}"]
    if topic:
        lines.append(f"> 주제: {topic}")
    lines.append("")
    lines.append(content)
    return "\n".join(lines)
